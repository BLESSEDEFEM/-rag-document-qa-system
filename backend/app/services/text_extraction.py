"""
Text extraction service.
Extracts text content from various document formats (PDF, DOCX, TXT).
"""


import logging
from typing import Dict, Any
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


async def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extract text from PDF file using PyPDF2.

    Args:
        file_path: Path to PDF file

    Returns:
        Dictionary with:
        - text: Extracted text content
        - page_count: Number of pages
        - success: True if extraction succeeded
        - error: Error message if failed

    Security: Handles corrupted PDFs gracefully (no crashes).
    """
    try:
        logger.info(f"Extracting text from PDF: {file_path}")

        with open(file_path, "rb") as file:
            # Create PDF Reader object
            pdf_reader = PyPDF2.PdfReader(file)

            # Get page count
            page_count = len(pdf_reader.pages)

            # Extract text from pages
            text_parts = []
            for page_num in range(page_count):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            # Combine all text
            full_text = "\n\n".join(text_parts)

            logger.info(f"PDF extraction successful: {page_count} pages, {len(full_text)} characters")

            return {
                "text": full_text,
                "page_count": page_count,
                "success": True,
                "error": None
            }

    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        return {
            "text": "",
            "page_count": None,
            "success": False,
            "error": str(e)
        }


async def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract text from DOCX file using python-docx.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with:
        - text: Extracted text content
        - page_count: None (DOCX doesn't have fixed pages)
        - success: True if extraction succeeded
        - error: Error message if failed

    Note: DOCX page count depends on formatting (not reliably extractable).
    """
    try:
        logger.info(f"Extracting text from DOCX: {file_path}")

        # Load document
        doc = Document(file_path)

        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Combine all texts
        full_text = "\n\n".join(text_parts)

        logger.info(f"DOCX extraction successful: {len(doc.paragraphs)} paragraphs, {len(full_text)} characters")

        return {
            "text": full_text,
            "page_count": None,
            "success": True,
            "error": None
        }
    except Exception as e:
        return {
            "text": "",
            "page_count": None,
            "success": False,
            "error": str(e)
        }


async def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """
    Extract text from TXT file (plain text reading).

    Args:
        file_path: Path to TXT file

    Returns:
        Dictionary with:
        - text: File contents
        - page_count: None (TXT doesn't have pages)
        - success: True if extraction succeeded
        - error: Error message if failed

    Handles various encodings (UTF-8, latin-1, etc.).
    """
    try:
        logger.info(f"Extracting text from TXT: {file_path}")

        # Try UTF-8 first (most common)
        try:
            with open(file_path, "r", encoding="utf_8") as file:
                text = file.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            logger.warning(f"UTF-8 decode failed for {file_path}, trying latin-1")
            with open(file_path, "r", encoding="latin-1") as file:
                text = file.read()

        logger.info(f"TXT extraction successful: {len(text)}")
        return {
            "text": text,
            "page_count": None,
            "success": True,
            "error": None
        }
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {file_path}: {e}")
        return {
            "text": "",
            "page_count": None,
            "success": False,
            "error": str(e)
        }


async def extract_text(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Extract text from document (auto-detects format).

    Args:
        file_path: Path to document file
        file_type: File extension (e.g., "pdf", "docx", "txt")

    Returns:
        Dictionary with extraction results (see individual functions)

    Routing function: Calls appropriate extractor based on file type.
    """
    file_type_lower = file_type.lower().replace(".", "")

    logger.info(f"Extracting text from {file_type_lower} file: {file_path}")

    if file_type_lower == "pdf":
        return await extract_text_from_pdf(file_path)
    elif file_type_lower == "docx":
        return await extract_text_from_docx(file_path)
    elif file_type_lower == "txt":
        return await extract_text_from_txt(file_path)
    else:
        logger.error(f"Unsupported file type: {file_type}")
        return {
            "text": "",
            "page_count": None,
            "success": False,
            "error": f"Unsupported file type: {file_type}"
        }
